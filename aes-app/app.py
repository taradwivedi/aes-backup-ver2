from flask import Flask, render_template, request, jsonify, send_file
import torch
import pickle
import numpy as np
import pandas as pd
import re
import textwrap

from utils import preprocess_text, extract_features, grammar_feedback
from docx import Document
from bs4 import BeautifulSoup
from docx.shared import RGBColor
from io import BytesIO
from reportlab.pdfgen import canvas
from reportlab.lib.colors import red, green, black


import tempfile
# ========== App Setup ==========
app = Flask(__name__)

import torch.nn as nn

class LSTMRegressor(nn.Module):
    def __init__(self, vocab_size, embed_dim, num_features):
        super(LSTMRegressor, self).__init__()
        self.embedding = nn.Embedding(vocab_size, embed_dim)
        self.lstm = nn.LSTM(embed_dim, hidden_size=64, batch_first=True)
        self.fc1 = nn.Linear(64 + num_features, 64)
        self.relu = nn.ReLU()
        self.out = nn.Linear(64, 1)

    def forward(self, text, num_feat):
        embedded = self.embedding(text)
        _, (hidden, _) = self.lstm(embedded)
        combined = torch.cat((hidden[-1], num_feat), dim=1)
        x = self.relu(self.fc1(combined))
        return self.out(x).squeeze()

# ========== Load Model & Preprocessors ==========

with open("tokenizer.pkl", "rb") as f:
    tokenizer = pickle.load(f)

vocab_size = 10000
embed_dim = 100
num_features = 27

model = LSTMRegressor(vocab_size, embed_dim, num_features)
model.load_state_dict(torch.load("lstm_model.pth", map_location=torch.device("cpu")))
model.eval()

MAX_LEN = 300

# ========== Good English Words ==========
GOOD_WORDS = [
    "moreover", "therefore", "consequently", "significant", "noteworthy", "intriguing", "persuasive", "elaborate",
    "nevertheless", "furthermore", "comprehensive", "profound", "insightful", "lucid", "relevant", "analytical",
    "structured", "articulate", "sophisticated", "convey", "illustrate", "substantiate", "compelling", "logical"
]

# ========== Routes ==========

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/predict', methods=['POST'])
def predict():
    data = request.form['essay']
    cleaned = preprocess_text(data)

# Server-side validation
    if not data.strip():
        return jsonify({'error': 'empty'})
    if not re.search(r'[a-zA-Z]', data):
        return jsonify({'error': 'invalid'})
    if len(cleaned.split()) < 50:
        return jsonify({'error': 'short'})
    # Grammar feedback
    highlighted_essay, feedback_list = grammar_feedback(data)

    # Extract features
    numeric_features = extract_features(cleaned)
    numeric_array = np.array(numeric_features).reshape(1, -1)

    # Tokenize & pad
    tokens = tokenizer.texts_to_sequences([cleaned])
    tokens_padded = np.zeros((1, MAX_LEN), dtype=np.int64)
    tokens_padded[0, :min(MAX_LEN, len(tokens[0]))] = tokens[0][:MAX_LEN]

    # Tensors
    text_tensor = torch.tensor(tokens_padded, dtype=torch.long)
    feat_tensor = torch.tensor(numeric_array, dtype=torch.float32)

    # Predict
    with torch.no_grad():
        prediction = model(text_tensor, feat_tensor).item()

    model_score = prediction * 10

    # Heuristic scoring
    grammar_error_count = len(feedback_list)
    heuristic_score = max(0, 10 - grammar_error_count * 0.4)
    grammar_percent = heuristic_score / 10

    # Extra Metrics
    word_count = len(cleaned.split())
    good_words_used = sum(1 for word in cleaned.lower().split() if word in GOOD_WORDS)

    # Adjustment
    adjustment = 0
    if grammar_error_count == 0:
        adjustment += 0.4
    elif grammar_error_count > 5:
        adjustment -= 0.2

    final_score = 0.75 * model_score + 0.25 * heuristic_score + adjustment

    # ----- Boost logic -----
    if word_count >= 200 and grammar_error_count == 0 and good_words_used >= 3:
        final_score = 10
    elif word_count >= 200 and grammar_percent >= 0.9 and good_words_used >= 2:
        final_score = 9
    elif word_count >= 200 and grammar_percent >= 0.9 and good_words_used >= 1:
        final_score = max(final_score, 8.5)

    # Round first, then clamp, and force perfect 10 if close enough
    rounded_score = round(final_score, 1)

    # If score is very close to 10, show it as 10.0
    if rounded_score >= 9.90:
        final_score = 10.0
    else:
        final_score = max(0, min(10, rounded_score))

    final_score_str = f"{final_score}/10"


    # Add feedback based on heuristics
    if grammar_error_count == 0:
        feedback_list.append("Excellent grammar usage.")
    if word_count >= 200:
        feedback_list.append("Essay has sufficient length and structure.")
    if good_words_used >= 2:
        feedback_list.append("Strong use of advanced vocabulary.")
    if good_words_used < 2:
        feedback_list.append("Use interesting words like moreover, therefore, consequently, significant to enhance your essay.")
    if word_count < 200:
        feedback_list.append("Add more valuable content to your essay.")
    elif grammar_error_count > 5:
        feedback_list.append("Consider correcting the spelling errors in your essay.")

    return jsonify({
        'score': final_score_str,
        'highlighted': highlighted_essay,
        'feedback': feedback_list
    })

# ========== NEW: Upload Endpoint ==========
@app.route('/upload', methods=['POST'])
def upload():
    uploaded_file = request.files.get('file')
    if not uploaded_file or not uploaded_file.filename.endswith(('.doc', '.docx')):
        return jsonify({'error': 'Invalid file format. Please upload a .doc or .docx file.'})

    try:
        doc = Document(uploaded_file)
        full_text = "\n".join([para.text for para in doc.paragraphs])
        return jsonify({'content': full_text})
    except Exception as e:
        return jsonify({'error': 'Failed to process the document.'})

# ========== NEW: Download Endpoint ==========


@app.route('/download', methods=['POST'])
def download():
    essay = request.form.get('essay', '')
    score = request.form.get('score', '')
    feedback = request.form.get('feedback', '')
    highlighted = request.form.get('highlighted', '')
    legend = request.form.get('legend', '')
    format = request.args.get('format', 'docx').lower()

    if format == 'docx':
        doc = Document()
        doc.add_heading("📝 Automated Essay Report", level=1)

        doc.add_paragraph("Essay:")
        doc.add_paragraph(essay)

        doc.add_heading("\n 1. 🏅 Final Score Evaluation:", level=2)
        doc.add_paragraph(score)

        doc.add_heading("\n 2. 🔍 Grammar Suggestions", level=2)
        doc.add_paragraph(feedback)

        doc.add_heading("\n 3. 🎯 Highlights:", level=2)

        # ✅ Only render the highlighted version (with red/green colors)
        soup = BeautifulSoup(highlighted, 'html.parser')
        para = doc.add_paragraph()
        for elem in soup.find_all(['span']):
            text = elem.get_text()
            style = elem.get('style', '')
            run = para.add_run(text)
            if 'color:red' in style:
                run.font.color.rgb = RGBColor(255, 0, 0)
            elif 'color:green' in style:
                run.font.color.rgb = RGBColor(0, 128, 0)

        doc.add_paragraph("\n\n(Generated by the AES System)")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return send_file(
            buffer,
            as_attachment=True,
            download_name="essay_report.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    elif format == 'pdf':
        buffer = BytesIO()
        c = canvas.Canvas(buffer, pagesize=(595, 842))  # A4 size in points

        def wrap_text(text, max_width=90):
            return textwrap.fill(text, width=max_width)

        y = 800
        line_height = 14

        def draw_text(title, content, bold=False, spacer=True):
            nonlocal y
            if spacer:
                y -= 10  # ✅ Add space before section
            content = str(content)
            if bold:
                c.setFont("Helvetica-Bold", 11)
                for line in wrap_text(title).split('\n'):
                    if y < 40:
                        c.showPage()
                        y = 800
                    c.setFillColor(black)
                    c.drawString(40, y, line)
                    y -= line_height

                # Now draw content in normal font
                c.setFont("Helvetica", 10)
                for line in wrap_text(content).split('\n'):
                    if y < 40:
                        c.showPage()
                        y = 800
                    c.setFillColor(black)
                    c.drawString(40, y, line)
                    y -= line_height
            else:
                c.setFont("Helvetica", 10)
                for line in wrap_text(f"{title}\n{content}").split('\n'):
                    if y < 40:
                        c.showPage()
                        y = 800
                    c.setFillColor(black)
                    c.drawString(40, y, line)
                    y -= line_height


        def draw_highlighted_text(html_text):
            nonlocal y
            c.setFont("Helvetica", 10)
            soup = BeautifulSoup(html_text, 'html.parser')

            for tag in soup.find_all('span', recursive=True):
                text = tag.get_text().strip()
                if not text:
                    continue

                color = black
                if tag.name == 'span':
                    style = tag.get('style', '')
                    if 'red' in style:
                        color = red
                    elif 'green' in style:
                        color = green

                for line in wrap_text(text).split('\n'):
                    if y < 40:
                        c.showPage()
                        y = 800
                    c.setFillColor(color)
                    c.drawString(40, y, line)
                    y -= line_height
                y-=4

        # Draw content
        draw_text("📝 Automated Essay Report", "", bold=True,spacer=False)
        
        draw_text("Essay:", essay)
        
        draw_text("1. Final Score Evaluation:", score,bold=True)
        
        draw_text("2. Grammar Suggestions", feedback,bold=True)
    
        draw_text("3. Highlights:", "", bold=True)
        draw_highlighted_text(highlighted)
        draw_text("Generated by:", "Automated Essay Scoring System")

        c.save()
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name="essay_report.pdf", mimetype='application/pdf')

    return jsonify({'error': 'Unsupported format'}), 400

@app.route('/about')
def about():
    return render_template('about.html')

@app.route('/goal')
def goal():
    return render_template('goal.html')

# ========== Main ==========
if __name__ == '__main__':
    app.run(debug=True)
